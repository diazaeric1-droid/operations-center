"""Convert a markdown AFE into a polished .docx with table-aware formatting."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

from .models import AFEDiagnosis


BRAND_COLOR = RGBColor(0x1F, 0x3A, 0x5F)
SECTION_RE = re.compile(r"^#{1,3}\s+(.+)$")
TABLE_LINE_RE = re.compile(r"^\|(.+)\|\s*$")


def _add_styled_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = BRAND_COLOR
    run.font.size = Pt(16 if level == 1 else 13 if level == 2 else 11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _add_markdown_table(doc: Document, lines: list[str]):
    """Render a contiguous block of markdown table lines into a Word table."""
    rows = []
    for line in lines:
        # Skip pure separator lines like |---|---|
        if re.match(r"^\|[\s\-:\|]+\|$", line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row_cells in enumerate(rows):
        for j in range(n_cols):
            cell_text = row_cells[j] if j < len(row_cells) else ""
            cell = table.rows[i].cells[j]
            # Render markdown **bold** spans instead of leaking literal asterisks.
            paragraph = cell.paragraphs[0]
            paragraph.text = ""
            for part in re.split(r"(\*\*[^*]+\*\*)", cell_text):
                if not part:
                    continue
                is_bold = part.startswith("**") and part.endswith("**")
                run = paragraph.add_run(part[2:-2] if is_bold else part)
                run.font.size = Pt(9)
                run.bold = bool(is_bold or i == 0)


def _add_paragraph(doc: Document, text: str):
    text = text.strip()
    if not text:
        return
    # Strip leading bullet marks; render as bullet if present
    bullet = text.startswith(("- ", "* ", "• "))
    if bullet:
        text = text[2:]
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    # Render **bold** segments
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    for run in p.runs:
        run.font.size = Pt(10)


def build_docx(markdown: str, output_path: str | Path, diagnosis: AFEDiagnosis) -> Path:
    """Build the AFE as a styled .docx file."""
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AUTHORIZATION FOR EXPENDITURE")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"{diagnosis.field}  ·  Well {diagnosis.well_id}  ·  {diagnosis.intervention.replace('_', ' ').title()}")
    sub_run.italic = True
    sub_run.font.size = Pt(11)

    doc.add_paragraph()  # spacer

    # Parse markdown line-by-line
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        section_m = SECTION_RE.match(line.strip())
        if section_m:
            level = len(line) - len(line.lstrip("#"))
            _add_styled_heading(doc, section_m.group(1).strip(), level=min(level, 3))
            i += 1
            continue

        if TABLE_LINE_RE.match(line):
            # collect contiguous table lines
            table_lines = []
            while i < len(lines) and TABLE_LINE_RE.match(lines[i]):
                table_lines.append(lines[i])
                i += 1
            _add_markdown_table(doc, table_lines)
            continue

        _add_paragraph(doc, line)
        i += 1

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
